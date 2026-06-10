from fastapi import FastAPI, Query, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pypinyin import pinyin, Style
import requests
import hashlib
import json
import re
import os
import io
import time
import sqlite3
import asyncio
from pathlib import Path
from datetime import datetime
from contextlib import asynccontextmanager

# ─── LIFESPAN (khởi tạo DB khi server start) ─────────────────────────────────
@asynccontextmanager
async def lifespan(app):
    init_db()
    yield

app = FastAPI(title="Dịch Trung-Việt API", version="3.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── PATHS ───────────────────────────────────────────────────────────────────
# /data là Render Disk — persistent storage, không mất khi restart
# Fallback về /tmp nếu chạy local hoặc chưa mount disk
DATA_DIR   = Path(os.getenv("DATA_DIR", "/data/translate_data"))
CACHE_DIR  = DATA_DIR / "cache"
AUDIO_DIR  = DATA_DIR / "audio"
DB_PATH    = DATA_DIR / "history.db"
# Font CJK cho PDF — trên Render Ubuntu cần cài thêm; fallback nếu không có
FONT_PATH  = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"

for d in [CACHE_DIR, AUDIO_DIR]:
    d.mkdir(parents=True, exist_ok=True)

CACHE_TTL  = 60 * 60 * 24 * 7   # 7 ngày

# ─── DATABASE ─────────────────────────────────────────────────────────────────
def get_conn():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_conn() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS history (
                id        INTEGER PRIMARY KEY AUTOINCREMENT,
                original  TEXT NOT NULL,
                translated TEXT NOT NULL,
                phonetic  TEXT,
                from_lang TEXT NOT NULL,
                to_lang   TEXT NOT NULL,
                provider  TEXT,
                created_at TEXT NOT NULL
            )
        """)
        conn.commit()

def save_history(original, translated, phonetic, from_lang, to_lang, provider):
    try:
        with get_conn() as conn:
            conn.execute("""
                INSERT INTO history (original, translated, phonetic, from_lang, to_lang, provider, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (original, translated, phonetic or "", from_lang, to_lang, provider,
                  datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
            conn.commit()
    except Exception:
        pass  # Không để lỗi DB phá request dịch

# ─── CACHE ────────────────────────────────────────────────────────────────────
def cache_key(text, from_lang, to_lang):
    return hashlib.md5(f"{text}|{from_lang}|{to_lang}".encode()).hexdigest()

def cache_get(key):
    f = CACHE_DIR / key
    if f.exists():
        try:
            data = json.loads(f.read_text())
            if time.time() - data["ts"] < CACHE_TTL:
                return data["value"]
        except Exception:
            pass
    return None

def cache_set(key, value):
    try:
        (CACHE_DIR / key).write_text(json.dumps({"ts": time.time(), "value": value}))
    except Exception:
        pass

# ─── CHUNKER ──────────────────────────────────────────────────────────────────
def chunk_text(text, max_len=1000):
    if len(text) <= max_len:
        return [text]
    chunks, current = [], ""
    for s in re.split(r'(?<=[。！？.!?\n])', text):
        if len(current) + len(s) <= max_len:
            current += s
        else:
            if current:
                chunks.append(current)
            if len(s) > max_len:
                for i in range(0, len(s), max_len):
                    chunks.append(s[i:i+max_len])
            else:
                current = s
    if current:
        chunks.append(current)
    return chunks

# ─── TRANSLATE PROVIDERS ──────────────────────────────────────────────────────
def translate_google(text, from_lang, to_lang):
    try:
        r = requests.get(
            "https://translate.googleapis.com/translate_a/single",
            params={"client": "gtx", "sl": from_lang, "tl": to_lang, "dt": "t", "q": text},
            timeout=8
        )
        r.raise_for_status()
        data = r.json()
        return "".join([i[0] for i in data[0] if i[0]]), ""
    except Exception:
        return "", ""

def translate_mymemory(text, from_lang, to_lang):
    try:
        lmap = {"zh": "zh-CN", "vi": "vi-VN"}
        r = requests.get(
            "https://api.mymemory.translated.net/get",
            params={"q": text, "langpair": f"{lmap.get(from_lang)}|{lmap.get(to_lang)}"},
            timeout=8
        )
        r.raise_for_status()
        data = r.json()
        if data.get("responseStatus") == 200:
            return data["responseData"]["translatedText"], ""
    except Exception:
        pass
    return "", ""

LIBRE_URL = os.getenv("LIBRE_TRANSLATE_URL", "")

def translate_libre(text, from_lang, to_lang):
    if not LIBRE_URL:
        return "", ""
    try:
        r = requests.post(
            f"{LIBRE_URL}/translate",
            json={"q": text, "source": from_lang, "target": to_lang, "format": "text"},
            timeout=10
        )
        r.raise_for_status()
        return r.json().get("translatedText", ""), ""
    except Exception:
        return "", ""

def translate_with_fallback(text, from_lang, to_lang):
    for fn, name in [(translate_google, "google"), (translate_mymemory, "mymemory"), (translate_libre, "libre")]:
        result, phonetic = fn(text, from_lang, to_lang)
        if result:
            return result, phonetic, name
    return "Không thể dịch lúc này", "", "none"

# ─── PINYIN ───────────────────────────────────────────────────────────────────
def get_pinyin(text):
    try:
        return ' '.join([p[0] for p in pinyin(text, style=Style.TONE, heteronym=False)])
    except Exception:
        return ""

# ─── CORE TRANSLATE ───────────────────────────────────────────────────────────
def do_translate(text, from_lang, to_lang, save_to_history=True):
    text = text.strip()
    if not text:
        raise HTTPException(400, "text không được rỗng")

    ck = cache_key(text, from_lang, to_lang)
    cached = cache_get(ck)
    if cached:
        cached["from_cache"] = True
        return cached

    chunks = chunk_text(text)
    parts, provider_used = [], "unknown"
    for chunk in chunks:
        translated, _, provider = translate_with_fallback(chunk, from_lang, to_lang)
        provider_used = provider
        parts.append(translated)

    full_translated = " ".join(parts)
    phonetic = get_pinyin(full_translated) if to_lang == "zh" else ""

    result = {
        "original": text,
        "translated": full_translated,
        "phonetic": phonetic,
        "provider": provider_used,
        "chunks": len(chunks),
        "from_cache": False,
    }
    cache_set(ck, result)
    if save_to_history:
        save_history(text, full_translated, phonetic, from_lang, to_lang, provider_used)
    return result

# ─── DETECT LANGUAGE ──────────────────────────────────────────────────────────
def detect_language(text):
    if any("\u4e00" <= c <= "\u9fff" for c in text):
        return "zh"
    if any(c in "àáảãạâầấẩẫậăằắẳẵặêềếểễệđ" for c in text.lower()):
        return "vi"
    return "vi"

# ─── EDGE TTS ─────────────────────────────────────────────────────────────────
# Giọng tốt nhất cho mỗi ngôn ngữ
TTS_VOICES = {
    "zh": "zh-CN-XiaoxiaoNeural",   # Giọng nữ Trung Quốc tự nhiên nhất
    "vi": "vi-VN-HoaiMyNeural",     # Giọng nữ Việt Nam
}

async def _generate_tts(text: str, lang: str) -> bytes:
    import edge_tts
    voice = TTS_VOICES.get(lang, "zh-CN-XiaoxiaoNeural")
    buf = io.BytesIO()
    communicate = edge_tts.Communicate(text, voice)
    async for chunk in communicate.stream():
        if chunk["type"] == "audio":
            buf.write(chunk["data"])
    return buf.getvalue()

# ─── EXPORT HELPERS ───────────────────────────────────────────────────────────
def _has_cjk_font():
    return Path(FONT_PATH).exists()

def export_pdf(rows: list) -> bytes:
    """Tạo PDF từ danh sách lịch sử. rows = list of sqlite3.Row"""
    from fpdf import FPDF, XPos, YPos

    class HistoryPDF(FPDF):
        def header(self):
            if _has_cjk_font():
                self.set_font("NotoSansCJK", size=10)
            else:
                self.set_font("Helvetica", size=10)
            self.set_text_color(120, 120, 120)
            self.cell(0, 8, "Lịch sử dịch — Translation History", align="C",
                      new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            self.ln(2)

        def footer(self):
            self.set_y(-15)
            if _has_cjk_font():
                self.set_font("NotoSansCJK", size=8)
            else:
                self.set_font("Helvetica", size=8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"Trang {self.page_no()}", align="C")

    pdf = HistoryPDF()
    if _has_cjk_font():
        pdf.add_font("NotoSansCJK", fname=FONT_PATH)

    pdf.add_page()
    font = "NotoSansCJK" if _has_cjk_font() else "Helvetica"

    # Tiêu đề
    pdf.set_font(font, size=16)
    pdf.set_text_color(30, 80, 160)
    pdf.cell(0, 12, "Lịch sử phiên dịch Trung - Việt", align="C",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font(font, size=9)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 6, f"Xuất lúc: {datetime.now().strftime('%d/%m/%Y %H:%M')}  •  Tổng: {len(rows)} mục",
             align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(6)

    # Header bảng
    COL = [65, 65, 40, 20]
    HEADERS = ["Văn bản gốc", "Bản dịch", "Phiên âm (Pinyin)", "Thời gian"]
    pdf.set_font(font, size=10)
    pdf.set_fill_color(30, 80, 160)
    pdf.set_text_color(255, 255, 255)
    for w, h in zip(COL, HEADERS):
        pdf.cell(w, 9, h, border=1, fill=True, new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.ln(9)

    # Rows
    pdf.set_font(font, size=9)
    for i, row in enumerate(rows):
        pdf.set_fill_color(245, 248, 255) if i % 2 == 0 else pdf.set_fill_color(255, 255, 255)
        pdf.set_text_color(30, 30, 30)
        h = 8
        y0 = pdf.get_y()
        # Check page break
        if y0 + h > pdf.h - 20:
            pdf.add_page()
            y0 = pdf.get_y()

        pdf.multi_cell(COL[0], h, str(row["original"])[:80], border=1,
                       fill=True, new_x=XPos.RIGHT, new_y=YPos.TOP, max_line_height=h)
        pdf.multi_cell(COL[1], h, str(row["translated"])[:80], border=1,
                       fill=True, new_x=XPos.RIGHT, new_y=YPos.TOP, max_line_height=h)
        pdf.multi_cell(COL[2], h, str(row["phonetic"] or "")[:50], border=1,
                       fill=True, new_x=XPos.RIGHT, new_y=YPos.TOP, max_line_height=h)
        ts = str(row["created_at"] or "")[:16]
        pdf.multi_cell(COL[3], h, ts, border=1,
                       fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT, max_line_height=h)

    return bytes(pdf.output())


def export_docx(rows: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margin
    for section in doc.sections:
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

    # Tiêu đề
    title = doc.add_heading("Lịch sử phiên dịch Trung - Việt", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"Xuất lúc: {datetime.now().strftime('%d/%m/%Y %H:%M')}  •  Tổng: {len(rows)} mục"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x78, 0x78, 0x78)
    sub.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Bảng
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header
    HDR = ["Văn bản gốc", "Bản dịch", "Phiên âm (Pinyin)", "Thời gian"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(HDR):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Màu nền header
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1E50A0")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    # Data rows
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = str(row["original"])
        cells[1].text = str(row["translated"])
        cells[2].text = str(row["phonetic"] or "")
        cells[3].text = str(row["created_at"] or "")[:16]
        # Màu nền xen kẽ
        fill_color = "EEF2FF" if i % 2 == 0 else "FFFFFF"
        for cell in cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill_color)
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"), "clear")
            tc_pr.append(shd)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/")
def root():
    return {"message": "Translation Server v3.0 — Trung↔Việt + TTS + History", "status": "ok"}

@app.get("/health")
def health():
    return {"status": "ok", "version": "3.0", "db": str(DB_PATH), "tts": "edge-tts"}


# ─── TRANSLATE ────────────────────────────────────────────────────────────────
@app.get("/translate")
def translate(
    text:      str = Query(..., min_length=1, max_length=10000),
    from_lang: str = Query(..., pattern="^(zh|vi)$"),
    to_lang:   str = Query(..., pattern="^(zh|vi)$"),
):
    if from_lang == to_lang:
        raise HTTPException(400, "from_lang và to_lang phải khác nhau")
    return do_translate(text, from_lang, to_lang)

@app.get("/auto-translate")
def auto_translate(text: str = Query(..., min_length=1, max_length=10000)):
    from_lang = detect_language(text)
    to_lang   = "vi" if from_lang == "zh" else "zh"
    result    = do_translate(text, from_lang, to_lang)
    result["detected_from"] = from_lang
    return result

@app.post("/translate-doc")
def translate_doc(body: dict):
    text      = body.get("text", "").strip()
    from_lang = body.get("from_lang", "vi")
    to_lang   = body.get("to_lang",   "zh")
    if not text:
        raise HTTPException(400, "Thiếu trường 'text'")
    if from_lang not in ("zh", "vi") or to_lang not in ("zh", "vi"):
        raise HTTPException(400, "Ngôn ngữ không hợp lệ")
    if from_lang == to_lang:
        raise HTTPException(400, "from_lang và to_lang phải khác nhau")
    return do_translate(text, from_lang, to_lang)

@app.get("/pinyin")
def get_pinyin_only(text: str = Query(..., min_length=1, max_length=5000)):
    return {"original": text, "phonetic": get_pinyin(text)}


# ─── TTS ──────────────────────────────────────────────────────────────────────
@app.get("/tts")
async def tts(
    text: str = Query(..., min_length=1, max_length=500),
    lang: str = Query(..., pattern="^(zh|vi)$"),
):
    """
    Trả về file audio MP3 của văn bản.
    Android dùng: MediaPlayer.setDataSource(url) hoặc download bytes.
    """
    # Cache audio theo hash
    audio_key = hashlib.md5(f"{text}|{lang}".encode()).hexdigest() + ".mp3"
    audio_file = AUDIO_DIR / audio_key

    if not audio_file.exists():
        try:
            audio_data = await _generate_tts(text, lang)
            if not audio_data:
                raise HTTPException(500, "Không tạo được audio")
            audio_file.write_bytes(audio_data)
        except Exception as e:
            raise HTTPException(500, f"TTS lỗi: {str(e)}")

    return StreamingResponse(
        io.BytesIO(audio_file.read_bytes()),
        media_type="audio/mpeg",
        headers={"Content-Disposition": f'attachment; filename="tts_{lang}.mp3"'}
    )


# ─── HISTORY ──────────────────────────────────────────────────────────────────
@app.get("/history")
def get_history(
    limit:  int = Query(50,  ge=1, le=500),
    offset: int = Query(0,   ge=0),
    lang:   str = Query("",  description="Lọc theo from_lang: zh hoặc vi"),
    search: str = Query("",  description="Tìm kiếm trong văn bản gốc"),
):
    """Lấy lịch sử dịch, hỗ trợ phân trang và lọc."""
    query  = "SELECT * FROM history WHERE 1=1"
    params = []
    if lang in ("zh", "vi"):
        query  += " AND from_lang = ?"
        params.append(lang)
    if search:
        query  += " AND (original LIKE ? OR translated LIKE ?)"
        params += [f"%{search}%", f"%{search}%"]
    query += " ORDER BY id DESC LIMIT ? OFFSET ?"
    params += [limit, offset]

    with get_conn() as conn:
        rows = conn.execute(query, params).fetchall()
        total = conn.execute(
            "SELECT COUNT(*) FROM history" + (
                " WHERE from_lang=?" if lang in ("zh","vi") else ""
            ), [lang] if lang in ("zh","vi") else []
        ).fetchone()[0]

    return {
        "total": total,
        "limit": limit,
        "offset": offset,
        "items": [dict(r) for r in rows]
    }

@app.delete("/history/{item_id}")
def delete_history_item(item_id: int):
    """Xóa một mục lịch sử."""
    with get_conn() as conn:
        conn.execute("DELETE FROM history WHERE id = ?", [item_id])
        conn.commit()
    return {"deleted": item_id}

@app.delete("/history")
def clear_history():
    """Xóa toàn bộ lịch sử."""
    with get_conn() as conn:
        count = conn.execute("SELECT COUNT(*) FROM history").fetchone()[0]
        conn.execute("DELETE FROM history")
        conn.commit()
    return {"deleted": count}


# ─── EXPORT ───────────────────────────────────────────────────────────────────
@app.get("/export/pdf")
def export_history_pdf(
    limit:  int = Query(200, ge=1, le=1000),
    search: str = Query(""),
):
    """Xuất lịch sử dịch ra file PDF."""
    query  = "SELECT * FROM history WHERE 1=1"
    params = []
    if search:
        query  += " AND (original LIKE ? OR translated LIKE ?)"
        params += [f"%{search}%", f"%{search}%"]
    query += " ORDER BY id DESC LIMIT ?"
    params.append(limit)

    with get_conn() as conn:
        rows = conn.execute(query, params).fetchall()

    if not rows:
        raise HTTPException(404, "Không có dữ liệu để xuất")

    pdf_bytes = export_pdf(rows)
    filename  = f"lich_su_dich_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

@app.get("/export/docx")
def export_history_docx(
    limit:  int = Query(200, ge=1, le=1000),
    search: str = Query(""),
):
    """Xuất lịch sử dịch ra file Word (.docx)."""
    query  = "SELECT * FROM history WHERE 1=1"
    params = []
    if search:
        query  += " AND (original LIKE ? OR translated LIKE ?)"
        params += [f"%{search}%", f"%{search}%"]
    query += " ORDER BY id DESC LIMIT ?"
    params.append(limit)

    with get_conn() as conn:
        rows = conn.execute(query, params).fetchall()

    if not rows:
        raise HTTPException(404, "Không có dữ liệu để xuất")

    docx_bytes = export_docx(rows)
    filename   = f"lich_su_dich_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ─── CACHE ────────────────────────────────────────────────────────────────────
@app.delete("/cache")
def clear_cache():
    count = sum(1 for f in CACHE_DIR.glob("*") if f.unlink() is None)
    return {"deleted": count}
